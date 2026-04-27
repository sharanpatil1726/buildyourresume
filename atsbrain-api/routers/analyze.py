from fastapi import APIRouter, Depends, HTTPException, status as http_status
from fastapi.responses import Response
from middleware import get_current_user, get_current_profile
from database import supabase_admin
from models import AnalyzeRequest, CoverLetterRequest, VerifyUnlockRequest
from services.anthropic_service import analyze_resume, generate_cover_letter, generate_optimized_resume
from services.razorpay_service import create_order, verify_payment_signature
from config import get_settings, SCAN_UNLOCK_AMOUNT
from loguru import logger
import io
import re

router = APIRouter()

# ── Template colour palette ───────────────────────────────────────────────────

_PALETTE: dict[str, str] = {
    "ats-classic": "111827",  "ats-minimal": "374151",  "ats-navy": "1e3a5f",
    "ats-clean":   "1e40af",  "ats-modern":  "7c3aed",  "ats-sharp": "0f172a",
    "ats-slate":   "475569",  "ats-teal":    "0f766e",  "ats-stone": "44403c",
    "ats-ink":     "312e81",  "std-violet":  "7c3aed",  "std-navy":  "1e3a5f",
    "std-forest":  "065f46",  "std-teal":    "0f766e",  "std-charcoal": "1f2937",
    "std-burgundy":"7f1d1d",  "std-cobalt":  "1d4ed8",  "std-gold":  "92400e",
    "std-slate":   "334155",  "std-indigo":  "4338ca",  "basic-clean": "374151",
    "basic-simple":"111827",  "basic-compact":"1f2937", "basic-academic":"1e3a5f",
    "basic-entry": "374151",  "mod-violet":  "7c3aed",  "mod-navy":  "1e3a5f",
    "mod-teal":    "0f766e",  "mod-slate":   "334155",  "mod-purple": "6d28d9",
    "mod-cyan":    "0891b2",  "mod-emerald": "059669",  "mod-graphite":"374151",
    "mod-rose":    "be185d",  "mod-amber":   "b45309",
}

# ── Resume text parser ────────────────────────────────────────────────────────

_SECTION_RE = re.compile(
    r"^(CONTACT|PROFILE|SUMMARY|OBJECTIVE|PROFESSIONAL|EXPERIENCE|WORK|EMPLOYMENT|"
    r"EDUCATION|SKILL|TECHNICAL|CERTIF|PROJECT|AWARD|ACHIEVEMENT|LANGUAGE|REFERENCE|"
    r"VOLUNTEER|INTEREST|ACTIVITY|PUBLICATION|RESEARCH|HONOR)",
    re.IGNORECASE,
)

def _is_heading(line: str) -> bool:
    t = line.strip()
    if not t or len(t) > 60:
        return False
    if t == t.upper() and len(t) >= 3 and re.search(r"[A-Z]", t):
        return True
    clean = re.sub(r"[:\-_|]", "", t).strip()
    return bool(_SECTION_RE.match(clean))

def _parse_resume(text: str) -> dict:
    lines = [l.strip() for l in text.split("\n") if l.strip()]
    if not lines:
        return {"name": "", "contact": [], "sections": []}

    name = lines[0]
    body_start = len(lines)
    for i in range(1, min(len(lines), 8)):
        if _is_heading(lines[i]):
            body_start = i
            break

    contact = lines[1:body_start]
    sections: list[dict] = []
    cur: dict | None = None
    for line in lines[body_start:]:
        if _is_heading(line):
            if cur:
                sections.append(cur)
            cur = {"title": line, "lines": []}
        elif cur is not None:
            cur["lines"].append(line)
    if cur:
        sections.append(cur)

    return {"name": name, "contact": contact, "sections": sections}

# ── DOCX generator ────────────────────────────────────────────────────────────

def _add_bottom_border(paragraph, hex_color: str) -> None:
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    pPr = paragraph._element.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "2")
    bottom.set(qn("w:color"), hex_color)
    pBdr.append(bottom)
    pPr.append(pBdr)

def _create_docx(text: str, template_id: str, target_role: str) -> bytes:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches

    primary_hex = _PALETTE.get(template_id, "111827")
    primary_rgb = tuple(int(primary_hex[i:i+2], 16) for i in (0, 2, 4))

    parsed = _parse_resume(text)
    doc = Document()

    for sec in doc.sections:
        sec.top_margin    = Inches(0.8)
        sec.bottom_margin = Inches(0.8)
        sec.left_margin   = Inches(0.85)
        sec.right_margin  = Inches(0.85)

    # Remove default empty paragraph
    for el in list(doc.paragraphs[0]._element.getparent()):
        pass
    if doc.paragraphs:
        doc.paragraphs[0]._element.getparent().remove(doc.paragraphs[0]._element)

    # Name
    np_ = doc.add_paragraph()
    np_.paragraph_format.space_after = Pt(3)
    nr = np_.add_run(parsed["name"])
    nr.bold = True
    nr.font.size = Pt(22)
    nr.font.color.rgb = RGBColor(*primary_rgb)

    # Contact
    if parsed["contact"]:
        cp = doc.add_paragraph(" | ".join(parsed["contact"]))
        cp.paragraph_format.space_after = Pt(6)
        if cp.runs:
            cp.runs[0].font.size = Pt(9)
            cp.runs[0].font.color.rgb = RGBColor(100, 100, 100)

    # Divider
    div_p = doc.add_paragraph()
    div_p.paragraph_format.space_after = Pt(2)
    _add_bottom_border(div_p, primary_hex)

    # Sections
    for section in parsed["sections"]:
        hp = doc.add_paragraph()
        hp.paragraph_format.space_before = Pt(10)
        hp.paragraph_format.space_after  = Pt(3)
        hr = hp.add_run(section["title"].upper())
        hr.bold = True
        hr.font.size = Pt(10)
        hr.font.color.rgb = RGBColor(*primary_rgb)
        _add_bottom_border(hp, primary_hex)

        for line in section["lines"]:
            if not line.strip():
                continue
            is_bullet = line.strip().startswith(("-", "•", "*"))
            content = line.strip().lstrip("-•* ").strip() if is_bullet else line.strip()
            try:
                p = doc.add_paragraph(style="List Bullet") if is_bullet else doc.add_paragraph()
                if is_bullet:
                    p.add_run(content)
                else:
                    p.add_run(line.strip())
            except Exception:
                p = doc.add_paragraph(content)
            p.paragraph_format.space_after = Pt(1)
            if p.runs:
                p.runs[0].font.size = Pt(10)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()

# ── PDF generator ─────────────────────────────────────────────────────────────

def _create_pdf(text: str, template_id: str) -> bytes:
    from fpdf import FPDF

    primary_hex = _PALETTE.get(template_id, "111827")
    pr = int(primary_hex[0:2], 16)
    pg = int(primary_hex[2:4], 16)
    pb = int(primary_hex[4:6], 16)

    parsed = _parse_resume(text)

    def safe(s: str) -> str:
        return s.encode("latin-1", errors="replace").decode("latin-1")

    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=18)
    pdf.add_page()
    pdf.set_margins(18, 18, 18)
    ew = pdf.w - 36

    # Name
    pdf.set_font("Helvetica", "B", 22)
    pdf.set_text_color(pr, pg, pb)
    pdf.multi_cell(ew, 10, safe(parsed.get("name", "")))

    # Contact
    contact = parsed.get("contact", [])
    if contact:
        pdf.set_font("Helvetica", "", 9)
        pdf.set_text_color(107, 114, 128)
        pdf.multi_cell(ew, 5, safe("  |  ".join(contact)))
    pdf.ln(3)

    # Divider
    y = pdf.get_y()
    pdf.set_draw_color(pr, pg, pb)
    pdf.set_line_width(0.5)
    pdf.line(18, y, pdf.w - 18, y)
    pdf.ln(6)

    # Sections
    for section in parsed.get("sections", []):
        if pdf.get_y() > pdf.h - 45:
            pdf.add_page()
        pdf.set_font("Helvetica", "B", 10)
        pdf.set_text_color(pr, pg, pb)
        pdf.multi_cell(ew, 6, safe(section["title"].upper()))
        y = pdf.get_y()
        pdf.set_draw_color(pr, pg, pb)
        pdf.set_line_width(0.3)
        pdf.line(18, y, pdf.w - 18, y)
        pdf.ln(2)

        pdf.set_font("Helvetica", "", 10)
        pdf.set_text_color(55, 65, 81)
        for line in section.get("lines", []):
            stripped = line.strip()
            if not stripped:
                continue
            is_bullet = stripped[0] in ("-", "*", "•")
            if is_bullet:
                content = stripped.lstrip("-*• ").strip()
                pdf.multi_cell(ew, 5, safe("  - " + content))
            else:
                pdf.multi_cell(ew, 5, safe(stripped))
        pdf.ln(5)

    return bytes(pdf.output())

settings = get_settings()

FREE_FIELDS = {"keyword_score", "format_score", "content_score", "readability_score",
               "top_skills", "strengths", "interview_likelihood", "ats_breakdown",
               "gap_analysis", "primary_role"}


def _free_result(result: dict) -> dict:
    return {k: result[k] for k in FREE_FIELDS if k in result}


@router.post("/")
async def run_analysis(
    body: AnalyzeRequest,
    user=Depends(get_current_user),
    profile=Depends(get_current_profile),
):
    if len(body.resume_text.strip()) < 100:
        raise HTTPException(status_code=400, detail="Resume text too short. Please upload your full resume.")

    try:
        result = analyze_resume(body.resume_text, body.target_role, body.experience_level)
    except Exception as e:
        logger.error(f"Claude analysis error: {e}")
        raise HTTPException(status_code=500, detail="Analysis failed. Please try again.")

    scan_id = None
    try:
        scan = supabase_admin.table("scans").insert({
            "user_id":          user.id,
            "resume_id":        body.resume_id,
            "target_role":      body.target_role,
            "experience_level": body.experience_level,
            "ats_score":        int(result["ats_score"]),
            "keyword_score":    int(result.get("keyword_score") or 0),
            "format_score":     int(result.get("format_score") or 0),
            "content_score":    int(result.get("content_score") or 0),
            "readability_score": int(result.get("readability_score") or 0),
            "result_json":      result,
            "is_unlocked":      False,
        }).execute()
        scan_id = scan.data[0]["id"]
        logger.info(f"Scan saved id={scan_id} score={result['ats_score']} role={body.target_role}")
    except Exception as db_err:
        logger.error(f"Failed to save scan: {db_err}")

    try:
        supabase_admin.table("profiles").update({
            "scans_used": profile["scans_used"] + 1
        }).eq("id", user.id).execute()
    except Exception as db_err:
        logger.warning(f"Failed to update scan count: {db_err}")

    # Pro users get the full result without needing per-scan unlock
    if profile.get("plan") == "pro":
        if scan_id:
            try:
                supabase_admin.table("scans").update({"is_unlocked": True}).eq("id", scan_id).execute()
            except Exception:
                pass
        return {"scan_id": scan_id, "is_unlocked": True, "result": result}

    return {"scan_id": scan_id, "is_unlocked": False, "result": _free_result(result)}


@router.get("/history")
async def scan_history(user=Depends(get_current_user)):
    result = (
        supabase_admin.table("scans")
        .select("id, target_role, ats_score, experience_level, created_at, is_unlocked")
        .eq("user_id", user.id)
        .order("created_at", desc=True)
        .limit(20)
        .execute()
    )
    return result.data


@router.get("/{scan_id}")
async def get_scan(scan_id: str, user=Depends(get_current_user), profile=Depends(get_current_profile)):
    result = (
        supabase_admin.table("scans")
        .select("*")
        .eq("id", scan_id)
        .eq("user_id", user.id)
        .single()
        .execute()
    )
    if not result.data:
        raise HTTPException(status_code=404, detail="Scan not found")

    scan = result.data
    full = scan.get("result_json") or {}
    # Pro users get full access without per-scan unlock
    is_unlocked = scan.get("is_unlocked") or profile.get("plan") == "pro"
    if is_unlocked:
        return {"scan_id": scan_id, "is_unlocked": True,
                "target_role": scan["target_role"], "experience_level": scan["experience_level"],
                "created_at": scan["created_at"], "result": full}

    return {"scan_id": scan_id, "is_unlocked": False,
            "target_role": scan["target_role"], "experience_level": scan["experience_level"],
            "created_at": scan["created_at"], "result": _free_result(full)}


@router.post("/{scan_id}/unlock")
async def create_unlock_order(scan_id: str, user=Depends(get_current_user)):
    scan = (
        supabase_admin.table("scans")
        .select("id, user_id, is_unlocked")
        .eq("id", scan_id)
        .eq("user_id", user.id)
        .single()
        .execute()
    )
    if not scan.data:
        raise HTTPException(status_code=404, detail="Scan not found")
    if scan.data.get("is_unlocked"):
        raise HTTPException(status_code=400, detail="Already unlocked")

    order = create_order(
        amount=SCAN_UNLOCK_AMOUNT,
        receipt=f"unlock_{scan_id[:8]}",
        notes={"scan_id": scan_id, "user_id": user.id},
    )
    return {"order_id": order["id"], "amount": SCAN_UNLOCK_AMOUNT, "currency": "INR", "key_id": settings.razorpay_key_id}


@router.post("/{scan_id}/unlock-verify")
async def verify_unlock(scan_id: str, body: VerifyUnlockRequest, user=Depends(get_current_user)):
    is_valid = verify_payment_signature(body.razorpay_order_id, body.razorpay_payment_id, body.razorpay_signature)
    if not is_valid:
        raise HTTPException(status_code=400, detail="Invalid payment signature")

    supabase_admin.table("scans").update({
        "is_unlocked": True,
        "unlock_payment_id": body.razorpay_payment_id,
    }).eq("id", scan_id).eq("user_id", user.id).execute()

    scan = supabase_admin.table("scans").select("result_json, target_role, experience_level").eq("id", scan_id).single().execute()
    logger.info(f"Scan {scan_id} unlocked by {user.id}")
    return {"is_unlocked": True, "result": scan.data.get("result_json", {})}


@router.get("/{scan_id}/optimized")
async def get_optimized_resume(scan_id: str, user=Depends(get_current_user), profile=Depends(get_current_profile)):
    if profile["plan"] != "pro":
        raise HTTPException(
            status_code=403,
            detail="Resume download requires Pro plan (₹299/month). Upgrade at /pricing.",
        )
    scan = (
        supabase_admin.table("scans")
        .select("optimized_resume, result_json, target_role")
        .eq("id", scan_id)
        .eq("user_id", user.id)
        .single()
        .execute()
    )
    if not scan.data:
        raise HTTPException(status_code=404, detail="Scan not found")

    if scan.data.get("optimized_resume"):
        return {"text": scan.data["optimized_resume"]}

    full = scan.data.get("result_json") or {}
    missing = full.get("missing_keywords", [])
    resume_text = full.get("_resume_text", "")
    target_role = scan.data["target_role"]

    try:
        optimized = generate_optimized_resume(resume_text, target_role, missing)
    except Exception as e:
        logger.error(f"Optimized resume generation failed: {e}")
        raise HTTPException(status_code=500, detail="Could not generate optimized resume")

    try:
        supabase_admin.table("scans").update({"optimized_resume": optimized}).eq("id", scan_id).execute()
    except Exception:
        pass

    return {"text": optimized}


@router.get("/{scan_id}/download/docx")
async def download_docx(
    scan_id: str,
    template_id: str = "ats-classic",
    user=Depends(get_current_user),
    profile=Depends(get_current_profile),
):
    if profile["plan"] != "pro":
        raise HTTPException(status_code=403, detail="Resume download requires Pro plan (₹299/month).")

    scan = (
        supabase_admin.table("scans")
        .select("optimized_resume, result_json, target_role")
        .eq("id", scan_id)
        .eq("user_id", user.id)
        .single()
        .execute()
    )
    if not scan.data:
        raise HTTPException(status_code=404, detail="Scan not found")

    text = scan.data.get("optimized_resume")
    if not text:
        full = scan.data.get("result_json") or {}
        missing = full.get("missing_keywords", [])
        resume_text = full.get("_resume_text", "")
        target_role = scan.data["target_role"]
        try:
            text = generate_optimized_resume(resume_text, target_role, missing)
            supabase_admin.table("scans").update({"optimized_resume": text}).eq("id", scan_id).execute()
        except Exception as e:
            logger.error(f"Optimized resume generation for DOCX failed: {e}")
            raise HTTPException(status_code=500, detail="Could not generate optimized resume")

    try:
        docx_bytes = _create_docx(text, template_id, scan.data["target_role"])
    except Exception as e:
        logger.error(f"DOCX creation failed: {e}")
        raise HTTPException(status_code=500, detail="Could not create DOCX file")

    safe_role = re.sub(r"[^\w\s-]", "", scan.data["target_role"]).strip().replace(" ", "_")
    filename = f"{safe_role}_{template_id}.docx"
    return Response(
        content=docx_bytes,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


@router.get("/{scan_id}/download/pdf")
async def download_pdf(
    scan_id: str,
    template_id: str = "ats-classic",
    user=Depends(get_current_user),
    profile=Depends(get_current_profile),
):
    if profile["plan"] != "pro":
        raise HTTPException(status_code=403, detail="Resume download requires Pro plan (₹299/month).")

    scan = (
        supabase_admin.table("scans")
        .select("optimized_resume, result_json, target_role")
        .eq("id", scan_id)
        .eq("user_id", user.id)
        .single()
        .execute()
    )
    if not scan.data:
        raise HTTPException(status_code=404, detail="Scan not found")

    text = scan.data.get("optimized_resume")
    if not text:
        full = scan.data.get("result_json") or {}
        missing = full.get("missing_keywords", [])
        resume_text = full.get("_resume_text", "")
        target_role = scan.data["target_role"]
        try:
            text = generate_optimized_resume(resume_text, target_role, missing)
            supabase_admin.table("scans").update({"optimized_resume": text}).eq("id", scan_id).execute()
        except Exception as e:
            logger.error(f"Optimized resume generation for PDF failed: {e}")
            raise HTTPException(status_code=500, detail="Could not generate optimized resume")

    try:
        pdf_bytes = _create_pdf(text, template_id)
    except Exception as e:
        logger.error(f"PDF creation failed: {e}")
        raise HTTPException(status_code=500, detail="Could not create PDF file")

    safe_role = re.sub(r"[^\w\s-]", "", scan.data["target_role"]).strip().replace(" ", "_")
    filename = f"{safe_role}_{template_id}.pdf"
    return Response(
        content=pdf_bytes,
        media_type="application/pdf",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


@router.post("/cover-letter")
async def cover_letter(
    body: CoverLetterRequest,
    user=Depends(get_current_user),
    profile=Depends(get_current_profile),
):
    if profile["plan"] == "free":
        raise HTTPException(status_code=403, detail="Cover letter generation requires Pro plan.")
    letter = generate_cover_letter(body.resume_text, body.target_role, body.company)
    return {"letter": letter}
